"""Resume parsing utilities for PDF and DOCX with enhanced hyperlink extraction."""
from fastapi import UploadFile
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from typing import List, Tuple
import tempfile

async def extract_hyperlinks_from_docx(doc: Document) -> List[Tuple[str, str]]:
    """
    Extract hyperlinks from DOCX document.
    Returns list of (text, url) tuples.
    """
    hyperlinks = []
    
    for rel in doc.part.rels.values():
        # Look for hyperlink relationships
        if "hyperlink" in rel.reltype:
            url = rel.target_ref
            # Try to find the link text in the document
            for paragraph in doc.paragraphs:
                if rel.rId in paragraph._element.xml:
                    text = paragraph.text.strip()
                    if text:
                        hyperlinks.append((text, url))
                        break
            
            # Also check in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if rel.rId in cell._element.xml:
                            text = cell.text.strip()
                            if text:
                                hyperlinks.append((text, url))
                                break
    
    return hyperlinks


async def extract_urls_from_text(text: str) -> List[Tuple[str, str]]:
    """
    Extract URLs from plain text.
    Returns list of (text, url) tuples.
    """
    import re
    
    urls = []
    # Match common profile URLs in text
    url_pattern = r'(https?://[^\s]+)|([A-Za-z]+\s+[A-Za-z]+\s+[A-Za-z]+\s+[A-Za-z]+)\s*(https?://[^\s]+)'
    
    matches = re.finditer(r'(https?://[^\s]+)', text)
    for match in matches:
        url = match.group(0)
        urls.append(("Profile", url))
    
    return urls


async def read_resume_text(file: UploadFile) -> Tuple[str, List[Tuple[str, str]]]:
    """
    Read and extract text from uploaded resume file (PDF, DOCX, or TXT).
    Also extracts hyperlinks when available.
    
    Args:
        file: Uploaded file from FastAPI form.
    
    Returns:
        tuple: (extracted_text, list_of_hyperlinks)
            - extracted_text: str - Plain text content
            - list_of_hyperlinks: List[Tuple[str, str]] - List of (label, url) tuples
    """
    content = await file.read()
    filename = file.filename.lower()
    hyperlinks = []
    text = ""
    
    if filename.endswith('.pdf'):
        # PDF extraction with pdfplumber
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            tmp.flush()
            with pdfplumber.open(tmp.name) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                    
                    # Try to extract URLs from page
                    urls_in_page = await extract_urls_from_text(page_text)
                    hyperlinks.extend(urls_in_page)
        
        return text, hyperlinks
    
    elif filename.endswith('.docx'):
        # DOCX extraction with python-docx
        import io
        doc = Document(io.BytesIO(content))
        
        # Extract text from paragraphs
        paragraphs_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs_text.append(paragraph.text)
                # Also check for hyperlinks in this paragraph
                for run in paragraph.runs:
                    # Check if run has hyperlink
                    rPr = run._element.get_or_add_rPr()
                    rStyle = rPr.find(qn('w:rStyle'))
                    if rStyle is not None and 'Hyperlink' in rStyle.get(qn('w:val')):
                        hyperlinks.append((run.text, ""))  # URL will be extracted separately
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs_text.append(cell.text)
        
        text = "\n".join(paragraphs_text)
        
        # Extract hyperlinks from relationships
        docx_hyperlinks = await extract_hyperlinks_from_docx(doc)
        hyperlinks.extend(docx_hyperlinks)
        
        return text, hyperlinks
    
    else:
        # Plain text or unknown format
        text = content.decode(errors="ignore")
        urls_in_text = await extract_urls_from_text(text)
        hyperlinks.extend(urls_in_text)
        return text, hyperlinks
